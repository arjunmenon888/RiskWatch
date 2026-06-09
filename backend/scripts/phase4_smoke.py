from __future__ import annotations

import shutil
import tempfile
import time
from pathlib import Path

import httpx
from docx import Document as DocxDocument
from sqlalchemy import delete

from app.db.session import SessionLocal
from app.models.document import Document
from app.models.game import Game
from app.models.topic import Topic
from app.models.user import User
from scripts.smoke_auth import create_test_user_headers

BASE_URL = "http://127.0.0.1:8003"


def main() -> None:
    stamp = int(time.time())
    temp_dir = Path(tempfile.mkdtemp(dir="C:/tmp"))
    docx_path = create_test_docx(temp_dir)
    owner_email = f"phase4owner{stamp}@example.com"
    other_email = f"phase4other{stamp}@example.com"
    game_id: int | None = None

    try:
        with httpx.Client(timeout=25) as client:
            owner_headers = signup(client, owner_email)
            other_headers = signup(client, other_email)
            game = client.post(
                f"{BASE_URL}/creator/games",
                headers=owner_headers,
                json={
                    "title": "Phase Four Topic Game",
                    "description": "topic smoke",
                    "category": "Safety",
                    "visibility": "private",
                },
            )
            game.raise_for_status()
            game_id = game.json()["id"]

            with docx_path.open("rb") as file:
                upload = client.post(
                    f"{BASE_URL}/api/games/{game_id}/documents/upload",
                    headers=owner_headers,
                    files={
                        "file": (
                            docx_path.name,
                            file,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                )
            upload.raise_for_status()
            print("upload", upload.status_code, upload.json()["status"], upload.json()["chunk_count"])

            extracted = client.post(
                f"{BASE_URL}/api/ai/extract-topics",
                headers=owner_headers,
                json={"game_id": game_id, "replace_existing": True},
            )
            extracted.raise_for_status()
            topics = extracted.json()
            print("extract", extracted.status_code, len(topics), topics[0]["difficulty"], topics[0]["source_chunk_ids"])

            first_topic_id = topics[0]["id"]
            update = client.patch(
                f"{BASE_URL}/api/games/{game_id}/topics/{first_topic_id}",
                headers=owner_headers,
                json={"title": "Edited Fire Safety Topic", "selected": False},
            )
            update.raise_for_status()
            print("update", update.status_code, update.json()["title"], update.json()["selected"])

            if len(topics) > 1:
                topic_ids = [topic["id"] for topic in reversed(topics)]
                reorder = client.post(
                    f"{BASE_URL}/api/games/{game_id}/topics/reorder",
                    headers=owner_headers,
                    json={"topic_ids": topic_ids},
                )
                reorder.raise_for_status()
                print("reorder", reorder.status_code, [topic["id"] for topic in reorder.json()])

            other_list = client.get(f"{BASE_URL}/api/games/{game_id}/topics", headers=other_headers)
            print("other_list", other_list.status_code)

            delete_topic = client.delete(f"{BASE_URL}/api/games/{game_id}/topics/{first_topic_id}", headers=owner_headers)
            print("delete_topic", delete_topic.status_code)
    finally:
        cleanup(owner_email, other_email, game_id)
        shutil.rmtree(temp_dir, ignore_errors=True)


def signup(client: httpx.Client, email: str) -> dict[str, str]:
    return create_test_user_headers(email, "Phase Four")


def create_test_docx(temp_dir: Path) -> Path:
    path = temp_dir / "phase4_topics.docx"
    document = DocxDocument()
    fire_text = (
        "Fire safety evacuation extinguisher alarm muster hazard response training "
        "teaches workers how to identify fire risks and respond safely. "
    )
    electrical_text = (
        "Electrical lockout tagout isolation voltage inspection equipment safety "
        "teaches workers how to prevent electrical incidents before maintenance. "
    )
    document.add_paragraph(fire_text * 24)
    document.add_paragraph(electrical_text * 24)
    document.save(path)
    return path


def cleanup(owner_email: str, other_email: str, game_id: int | None) -> None:
    with SessionLocal() as db:
        if game_id is not None:
            db.execute(delete(Topic).where(Topic.game_id == game_id))
            db.execute(delete(Document).where(Document.game_id == game_id))
            db.execute(delete(Game).where(Game.id == game_id))
            shutil.rmtree(Path("storage") / "games" / str(game_id), ignore_errors=True)
        db.execute(delete(User).where(User.email.in_([owner_email, other_email])))
        db.commit()


if __name__ == "__main__":
    main()
